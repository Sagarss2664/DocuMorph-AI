import streamlit as st

# MUST be the very first Streamlit command
st.set_page_config(
    page_title="DocuMorph AI Pro",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/sagarss2664/DocuMorph-AI',
        'Report a bug': "https://github.com/sagarss2664/DocuMorph-AI/issues",
        'About': "# DocuMorph AI - Intelligent Document Formatter"
    }
)

# Now import other libraries
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import json
import pdfplumber
from docx2python import docx2python
from textblob import TextBlob
from PIL import Image
import nltk
import subprocess
import sys
import pandas as pd
from difflib import SequenceMatcher
import re

# -------------------- Setup NLTK Data --------------------
@st.cache_resource
def setup_nltk():
    try:
        nltk.data.find('tokenizers/punkt')
        return True
    except LookupError:
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            nltk.download('brown', quiet=True)
            return True
        except Exception as e:
            st.error(f"NLTK data download failed: {str(e)}")
            st.warning("""
            **Manual fix required:**
            1. Run this command locally:
            ```bash
            python -m textblob.download_corpora
            ```
            2. Then redeploy your app
            """)
            return False

if not setup_nltk():
    st.stop()

# -------------------- Enhanced Grammar Tools --------------------
def get_closest_match(word):
    """Get closest dictionary match for misspelled words"""
    from textblob import Word
    suggestions = Word(word).spellcheck()
    if suggestions:
        return suggestions[0][0]  # Return the most likely correction
    return word

def improved_grammar_check(text):
    """Enhanced grammar checking with better suggestions"""
    try:
        # First get TextBlob's suggestions
        blob = TextBlob(text)
        corrected = str(blob.correct())
        
        # Use difflib to find changes between original and corrected
        matcher = SequenceMatcher(None, text, corrected)
        issues = []
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                original = text[i1:i2]
                suggested = corrected[j1:j2]
                
                # Get context (50 chars before and after)
                start = max(0, i1 - 20)
                end = min(len(text), i2 + 20)
                context = text[start:end]
                
                issues.append({
                    "type": "Grammar/Spelling",
                    "original": original,
                    "suggestion": suggested,
                    "context": context
                })
        
        # Additional checks for common errors
        common_errors = {
            r"\btheir\b": ["there", "they're"],
            r"\byour\b": ["you're"],
            r"\bits\b": ["it's"],
            r"\baffect\b": ["effect"],
            r"\bthen\b": ["than"]
        }
        
        for pattern, alternatives in common_errors.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                for alt in alternatives:
                    if re.search(alt, text, re.IGNORECASE):
                        issues.append({
                            "type": "Common Error",
                            "original": match.group(),
                            "suggestion": f"Possible confusion with '{alt}'",
                            "context": text[max(0, match.start()-20):match.end()+20]
                        })
        
        return issues[:100]  # Limit to 100 issues
    except Exception as e:
        st.error(f"Grammar check error: {str(e)}")
        return []

# -------------------- Text Processing Tools --------------------
def extract_text_from_file(uploaded_file):
    """Extract text from PDF, DOC, or DOCX files with better formatting"""
    try:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                return text.strip()
        
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  "application/msword"]:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(uploaded_file.getvalue())
                docx_content = docx2python(tmp.name)
                text = docx_content.text
                os.unlink(tmp.name)
                return text
    except Exception as e:
        st.error(f"Text extraction error: {str(e)}")
        return ""

# -------------------- DocuMorph Engine --------------------
class DocuMorphEngine:
    def __init__(self, docx_file=None):
        self.document = Document(docx_file) if docx_file else Document()

    def set_font(self, font_name, font_size):
        for para in self.document.paragraphs:
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    def set_line_spacing(self, spacing):
        for para in self.document.paragraphs:
            para.paragraph_format.line_spacing = spacing

    def set_alignment(self, alignment):
        align_map = {
            "Left": WD_ALIGN_PARAGRAPH.LEFT,
            "Center": WD_ALIGN_PARAGRAPH.CENTER,
            "Right": WD_ALIGN_PARAGRAPH.RIGHT,
            "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        for para in self.document.paragraphs:
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def set_margins(self, top, bottom, left, right):
        sec = self.document.sections[0]
        sec.top_margin = Inches(top)
        sec.bottom_margin = Inches(bottom)
        sec.left_margin = Inches(left)
        sec.right_margin = Inches(right)

    def add_logo(self, image, width, height):
        hdr = self.document.sections[0].header
        if not hdr.paragraphs:
            hdr.add_paragraph()
        run = hdr.paragraphs[0].add_run()
        run.add_picture(image, width=Inches(width), height=Inches(height))

    def set_header_footer(self, h_text, f_text, size, align):
        align_map = {
            "Left": WD_ALIGN_PARAGRAPH.LEFT,
            "Center": WD_ALIGN_PARAGRAPH.CENTER,
            "Right": WD_ALIGN_PARAGRAPH.RIGHT
        }
        for sec in self.document.sections:
            # Header
            if not sec.header.paragraphs:
                sec.header.add_paragraph()
            h_para = sec.header.paragraphs[0]
            h_para.text = h_text
            if h_para.runs:
                h_para.runs[0].font.size = Pt(size)
            h_para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            
            # Footer
            if not sec.footer.paragraphs:
                sec.footer.add_paragraph()
            f_para = sec.footer.paragraphs[0]
            f_para.text = f_text
            if f_para.runs:
                f_para.runs[0].font.size = Pt(size)
            f_para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    def add_section_title(self, title):
        self.document.add_heading(title, level=1)

    def add_bullet_list(self, items):
        for item in items:
            self.document.add_paragraph(item, style="List Bullet")

    def add_figure(self, image, w, h, caption="", pos="Below"):
        if pos == "Above" and caption:
            self.document.add_paragraph(caption, style="Caption")
        p = self.document.add_paragraph()
        run = p.add_run()
        run.add_picture(image, width=Inches(w), height=Inches(h))
        if pos == "Below" and caption:
            self.document.add_paragraph(caption, style="Caption")

    def save(self, path):
        self.document.save(path)

# -------------------- Template Manager --------------------
TEMPLATE_DIR = "templates"
os.makedirs(TEMPLATE_DIR, exist_ok=True)

def list_templates():
    return [f[:-5] for f in os.listdir(TEMPLATE_DIR) if f.endswith('.json')]

def load_template(name):
    path = os.path.join(TEMPLATE_DIR, f"{name}.json")
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return None

def save_template(name, cfg):
    with open(os.path.join(TEMPLATE_DIR, f"{name}.json"), "w") as f:
        json.dump(cfg, f, indent=4)

def delete_template(name):
    path = os.path.join(TEMPLATE_DIR, f"{name}.json")
    if os.path.exists(path):
        os.remove(path)

# -------------------- Streamlit UI --------------------
# Custom CSS
st.markdown("""
<style>
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 8px 16px;
        border-radius: 4px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #f0f2f6;
    }
    .grammar-error {
        color: #d32f2f;
        background-color: #ffebee;
        padding: 2px 4px;
        border-radius: 3px;
    }
    .grammar-suggestion {
        color: #388e3c;
        background-color: #e8f5e9;
        padding: 2px 4px;
        border-radius: 3px;
    }
    .error-type {
        font-size: 0.85em;
        color: #616161;
    }
    .stButton>button {
        transition: all 0.2s ease;
    }
    .stButton>button:hover {
        transform: translateY(-1px);
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
    }
    .stDownloadButton>button {
        background-color: #4e79a7 !important;
        color: white !important;
    }
    .extracted-text {
        border: 1px solid #e1e4e8;
        border-radius: 6px;
        padding: 16px;
        margin-bottom: 16px;
        background-color: #f6f8fa;
        max-height: 300px;
        overflow-y: auto;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar - Template Manager
with st.sidebar:
    st.header("💾 Template Manager")
    templates = list_templates()
    selected_template = st.selectbox("Load Template", ["<none>"] + templates)
    
    if selected_template != "<none>":
        config = load_template(selected_template)
    else:
        config = {
            "font_name": "Times New Roman",
            "font_size": 12,
            "line_spacing": 1.15,
            "alignment": "Left",
            "margins": [1.0, 1.0, 1.0, 1.0],
            "header_text": "",
            "footer_text": "",
            "hf_size": 10,
            "hf_align": "Center",
            "logo_width": 1.0,
            "logo_height": 1.0
        }
    
    st.divider()
    new_template_name = st.text_input("Save Current Settings As")
    if st.button("💾 Save Template", use_container_width=True):
        if new_template_name:
            current_settings = {
                "font_name": st.session_state.get("font_name", "Times New Roman"),
                "font_size": st.session_state.get("font_size", 12),
                "line_spacing": st.session_state.get("line_spacing", 1.15),
                "alignment": st.session_state.get("alignment", "Left"),
                "margins": [
                    st.session_state.get("margin_top", 1.0),
                    st.session_state.get("margin_bottom", 1.0),
                    st.session_state.get("margin_left", 1.0),
                    st.session_state.get("margin_right", 1.0)
                ],
                "header_text": st.session_state.get("header_text", ""),
                "footer_text": st.session_state.get("footer_text", ""),
                "hf_size": st.session_state.get("hf_size", 10),
                "hf_align": st.session_state.get("hf_align", "Center"),
                "logo_width": st.session_state.get("logo_width", 1.0),
                "logo_height": st.session_state.get("logo_height", 1.0)
            }
            save_template(new_template_name, current_settings)
            st.success(f"Template '{new_template_name}' saved!")
        else:
            st.error("Please enter a template name")
    
    if selected_template != "<none>":
        if st.button("🗑 Delete Template", use_container_width=True):
            delete_template(selected_template)
            st.experimental_rerun()

# Main App
st.title("📄 DocuMorph AI Pro - Intelligent Document Formatter")

tab1, tab2, tab3, tab4 = st.tabs(["Formatting", "Content", "Grammar Check", "Export"])

# Tab 1: Formatting
with tab1:
    st.subheader("🎨 Document Styling")
    col1, col2 = st.columns(2)
    with col1:
        font_name = st.selectbox(
            "Font", 
            ["Times New Roman", "Arial", "Calibri", "Georgia"],
            index=["Times New Roman", "Arial", "Calibri", "Georgia"].index(config.get("font_name", "Times New Roman")),
            key="font_name"
        )
        font_size = st.slider(
            "Font Size", 8, 24, 
            config.get("font_size", 12),
            key="font_size"
        )
        line_spacing = st.slider(
            "Line Spacing", 1.0, 2.0, 
            config.get("line_spacing", 1.15), 0.05,
            key="line_spacing"
        )
    with col2:
        alignment = st.selectbox(
            "Alignment",
            ["Left", "Center", "Right", "Justify"],
            index=["Left", "Center", "Right", "Justify"].index(config.get("alignment", "Left")),
            key="alignment"
        )
        st.write("Margins (inches):")
        margins = st.columns(4)
        margin_labels = ["Top", "Bottom", "Left", "Right"]
        for i, label in enumerate(margin_labels):
            with margins[i]:
                st.number_input(
                    label,
                    0.1, 3.0,
                    config.get("margins", [1.0]*4)[i], 0.1,
                    key=f"margin_{label.lower()}"
                )

# Tab 2: Content
with tab2:
    st.subheader("📝 Document Content")
    col1, col2 = st.columns(2)
    with col1:
        logo = st.file_uploader(
            "Upload Logo (PNG/JPG)", 
            type=["png", "jpg", "jpeg"],
            key="logo"
        )
        if logo:
            logo_width = st.slider(
                "Logo Width (inches)", 0.5, 4.0,
                config.get("logo_width", 1.0), 0.1,
                key="logo_width"
            )
            logo_height = st.slider(
                "Logo Height (inches)", 0.5, 4.0,
                config.get("logo_height", 1.0), 0.1,
                key="logo_height"
            )
    with col2:
        header_text = st.text_input(
            "Header Text",
            config.get("header_text", ""),
            key="header_text"
        )
        footer_text = st.text_input(
            "Footer Text",
            config.get("footer_text", ""),
            key="footer_text"
        )
        hf_size = st.slider(
            "Header/Footer Font Size", 8, 20,
            config.get("hf_size", 10),
            key="hf_size"
        )
        hf_align = st.selectbox(
            "Header/Footer Alignment",
            ["Left", "Center", "Right"],
            index=["Left", "Center", "Right"].index(config.get("hf_align", "Center")),
            key="hf_align"
        )
    
    st.subheader("📑 Sections & Media")
    section_title = st.text_input("Section Title", key="section_title")
    bullets_input = st.text_area(
        "Bullet Points (one per line)", 
        height=100,
        key="bullets"
    )
    figure = st.file_uploader(
        "Add Figure (PNG/JPG)", 
        type=["png", "jpg", "jpeg"],
        key="figure"
    )
    if figure:
        fig_col1, fig_col2 = st.columns(2)
        with fig_col1:
            fig_width = st.slider("Width (inches)", 1.0, 6.0, 4.0, 0.1, key="fig_width")
        with fig_col2:
            fig_height = st.slider("Height (inches)", 1.0, 6.0, 3.0, 0.1, key="fig_height")
        caption = st.text_input("Caption", key="caption")
        caption_pos = st.radio(
            "Caption Position",
            ["Above", "Below"],
            horizontal=True,
            key="caption_pos"
        )

# Tab 3: Grammar Check
with tab3:
    st.subheader("🔍 Grammar & Spell Check")
    
    check_option = st.radio(
        "Check content from:",
        ["Upload Document", "Enter Text Directly"],
        horizontal=True,
        key="check_source"
    )
    
    text_to_check = ""
    extracted_text = ""
    
    if check_option == "Upload Document":
        grammar_file = st.file_uploader(
            "Upload Document (PDF/DOCX/TXT)", 
            type=["pdf", "docx", "txt"],
            key="grammar_file"
        )
        if grammar_file:
            with st.spinner("Extracting text..."):
                extracted_text = extract_text_from_file(grammar_file)
                if extracted_text:
                    st.subheader("Extracted Text Preview")
                    st.markdown(f'<div class="extracted-text">{extracted_text[:5000]}</div>', unsafe_allow_html=True)
                    text_to_check = extracted_text
    else:
        text_to_check = st.text_area(
            "Enter text to analyze",
            height=200,
            key="direct_text"
        )
    
    if text_to_check and st.button("Run Advanced Grammar Check", use_container_width=True):
        with st.spinner("Analyzing content..."):
            issues = improved_grammar_check(text_to_check[:10000])  # Limit to first 10k chars
            
            if not issues:
                st.success("✅ No grammar or spelling issues found!")
            else:
                st.warning(f"⚠️ Found {len(issues)} potential issues:")
                
                # Group issues by type
                issue_types = {}
                for issue in issues:
                    if issue['type'] not in issue_types:
                        issue_types[issue['type']] = []
                    issue_types[issue['type']].append(issue)
                
                # Display organized results
                for issue_type, issues in issue_types.items():
                    with st.expander(f"{issue_type} ({len(issues)} issues)", expanded=True):
                        for i, issue in enumerate(issues[:20]):  # Show first 20 per type
                            st.markdown(f"""
                            **{i+1}. {issue['type']}**  
                            <span class="error-type">Context: {issue['context']}</span>  
                            <span class="grammar-error">Original:</span> `{issue['original']}`  
                            <span class="grammar-suggestion">Suggestion:</span> `{issue['suggestion']}`
                            """, unsafe_allow_html=True)
                        if len(issues) > 20:
                            st.info(f"Showing first 20 of {len(issues)} {issue_type.lower()} issues")

# Tab 4: Export
with tab4:
    st.subheader("📤 Export Document")
    doc_file = st.file_uploader(
        "Upload DOCX to Format",
        type=["docx"],
        key="doc_file"
    )
    
    if st.button("Generate Formatted Document", use_container_width=True, type="primary"):
        if not doc_file:
            st.error("Please upload a DOCX file first")
        else:
            with st.spinner("Formatting document..."):
                try:
                    # Initialize engine
                    engine = DocuMorphEngine(doc_file)
                    
                    # Apply formatting
                    engine.set_font(
                        st.session_state.font_name,
                        st.session_state.font_size
                    )
                    engine.set_line_spacing(st.session_state.line_spacing)
                    engine.set_alignment(st.session_state.alignment)
                    engine.set_margins(
                        st.session_state.margin_top,
                        st.session_state.margin_bottom,
                        st.session_state.margin_left,
                        st.session_state.margin_right
                    )
                    
                    # Add logo if uploaded
                    if 'logo' in st.session_state and st.session_state.logo:
                        st.session_state.logo.seek(0)
                        engine.add_logo(
                            st.session_state.logo,
                            st.session_state.logo_width,
                            st.session_state.logo_height
                        )
                    
                    # Add header/footer
                    engine.set_header_footer(
                        st.session_state.header_text,
                        st.session_state.footer_text,
                        st.session_state.hf_size,
                        st.session_state.hf_align
                    )
                    
                    # Add content
                    if st.session_state.section_title.strip():
                        engine.add_section_title(st.session_state.section_title.strip())
                    
                    if st.session_state.bullets.strip():
                        bullets = [b.strip() for b in st.session_state.bullets.split("\n") if b.strip()]
                        engine.add_bullet_list(bullets)
                    
                    if 'figure' in st.session_state and st.session_state.figure:
                        st.session_state.figure.seek(0)
                        engine.add_figure(
                            st.session_state.figure,
                            st.session_state.fig_width,
                            st.session_state.fig_height,
                            st.session_state.caption,
                            st.session_state.caption_pos
                        )
                    
                    # Save and offer download
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        engine.save(tmp.name)
                        with open(tmp.name, "rb") as f:
                            st.download_button(
                                "⬇ Download Formatted Document",
                                f.read(),
                                "formatted.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    os.unlink(tmp.name)
                    
                except Exception as e:
                    st.error(f"Error generating document: {str(e)}")